# chunking.py
"""
Step 1: Extract, normalize, and chunk documents into JSON / JSONL.
You already have chunks.json/chunks.jsonl; re-run this only when adding new docs.

Usage:
    python -m chunking --in ks-*.pdf ks-*.docx --out data/chunks.json --jsonl
"""

import argparse, json, re, uuid, zipfile
from pathlib import Path
from typing import List, Dict, Any, Tuple

from .rag_config import DATA_DIR, CHUNKS_JSON, CHUNKS_JSONL

# Optional dependencies
try:
    import PyPDF2
    HAVE_PYPDF2 = True
except Exception:
    HAVE_PYPDF2 = False

try:
    import docx  # python-docx
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

def clean_text(s: str) -> str:
    s = s.replace("\u00A0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\r", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def chunk_words(text: str, target_words: int = 800, overlap_words: int = 100) -> List[str]:
    words = re.findall(r"\S+", text)
    chunks, i, n = [], 0, len(words)
    while i < n:
        j = min(i + target_words, n)
        chunk = " ".join(words[i:j])
        if chunk.strip():
            chunks.append(chunk.strip())
        if j >= n:
            break
        i = max(j - overlap_words, i + 1)
    return chunks

def infer_session_from_filename(name: str) -> str:
    stem = Path(name).stem
    if stem.startswith("ks-"):
        stem = stem[3:]
    return stem.replace("-", " ").title()

def extract_pdf(path: Path, max_pages: int | None = None) -> List[Tuple[int, str]]:
    pages = []
    if HAVE_PYPDF2:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            total = len(reader.pages)
            lim = min(total, max_pages) if max_pages else total
            for i in range(lim):
                t = (reader.pages[i].extract_text() or "")
                pages.append((i + 1, clean_text(t)))
    else:
        # best-effort fallback: read bytes and try a naive decode
        t = path.read_bytes().decode("latin-1", errors="ignore")
        pages = [(1, clean_text(t[:8000]))]
    return pages

def extract_docx(path: Path) -> List[str]:
    blocks: List[str] = []
    if HAVE_DOCX:
        d = docx.Document(str(path))
        for p in d.paragraphs:
            txt = clean_text(p.text)
            if txt:
                blocks.append(txt)
        for table in d.tables:
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                row_txt = " | ".join([c for c in cells if c])
                if row_txt:
                    blocks.append(row_txt)
    else:
        with zipfile.ZipFile(str(path)) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        xml = re.sub(r"<.*?>", "", xml)
        for line in xml.splitlines():
            line = clean_text(line)
            if line:
                blocks.append(line)
    return blocks

def make_chunks_for_pdf(path: Path) -> List[Dict[str, Any]]:
    pages = extract_pdf(path)
    text = "\n\n".join(f"[Page {p}]\n{t}" for p, t in pages if t)
    out = []
    for ch in chunk_words(text):
        out.append({
            "id": str(uuid.uuid4()),
            "text": ch,
            "meta": {
                "source": path.name,
                "type": "pdf",
                "page_start": pages[0][0] if pages else None,
                "page_end": pages[-1][0] if pages else None,
                "section": "",
                "session": infer_session_from_filename(path.name),
                "speaker": None,
            }
        })
    return out

def make_chunks_for_docx(path: Path) -> List[Dict[str, Any]]:
    blocks = extract_docx(path)
    text = "\n\n".join(blocks)
    out = []
    for ch in chunk_words(text):
        out.append({
            "id": str(uuid.uuid4()),
            "text": ch,
            "meta": {
                "source": path.name,
                "type": "docx",
                "section": "",
                "session": infer_session_from_filename(path.name),
                "speaker": None,
            }
        })
    return out

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="inputs", nargs="+", required=True, help="Input files (pdf/docx)")
    ap.add_argument("--out", default=str(CHUNKS_JSON), help="Output JSON path (array)")
    ap.add_argument("--jsonl", action="store_true", help="Also write JSONL alongside JSON")
    args = ap.parse_args()

    files = [Path(p) for p in args.inputs]
    chunks: List[Dict[str, Any]] = []
    for p in files:
        if not p.exists():
            print(f"Skip missing: {p}")
            continue
        if p.suffix.lower() == ".pdf":
            chunks.extend(make_chunks_for_pdf(p))
        elif p.suffix.lower() == ".docx":
            chunks.extend(make_chunks_for_docx(p))
        else:
            print(f"Skip unsupported: {p}")

    out_json = Path(args.out)
    out_json.parent.mkdir(parents=True, exist_ok=True)
    out_json.write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {out_json} with {len(chunks)} chunks.")

    if args.jsonl:
        out_jsonl = CHUNKS_JSONL if out_json == CHUNKS_JSON else out_json.with_suffix(".jsonl")
        with open(out_jsonl, "w", encoding="utf-8") as f:
            for obj in chunks:
                f.write(json.dumps(obj, ensure_ascii=False) + "\n")
        print(f"Wrote {out_jsonl}")

if __name__ == "__main__":
    main()
